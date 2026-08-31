"""Word output for the IC committee.

The committee attaches this to the quarterly board deck, so it has to be a
document rather than a data dump: the table first, the caveats close enough to
the numbers to be read with them, and the provenance far enough back to not be
in the way.

The design constraint is the same one that governs everything else here. A cell
that renders as blank in the Markdown output must render as blank *with its
reason* in Word, because this is the version that leaves the building. A Word
document that quietly looks more complete than the underlying evidence would
undo the entire point.
"""

from __future__ import annotations

import logging
from pathlib import Path

from ..config import ALL_METRICS, METRIC_LABELS
from ..core.models import Confidence
from ..pipeline import BenchmarkRun
from .comparison import compare
from .coverage import coverage_rows, FILLED
from .table import APEX_COLUMN, build_cells, format_value

log = logging.getLogger(__name__)

CONFIDENCE_MARK = {
    Confidence.HIGH: "High",
    Confidence.MEDIUM: "Med",
    Confidence.LOW: "Low",
}


class WordUnavailable(RuntimeError):
    pass


def _require_docx():
    try:
        import docx  # noqa: PLC0415

        return docx
    except ImportError as exc:  # pragma: no cover - environment dependent
        raise WordUnavailable(
            "python-docx is not installed. Install it with "
            "`pip install python-docx` (it is in requirements.txt), or use the "
            "Markdown outputs instead."
        ) from exc


def _cell_text(cell) -> str:
    if cell.value is None:
        reason = cell.reason.label if cell.reason else "not available"
        return f"— {reason}"
    text = format_value(cell.value, cell.unit)
    if cell.confidence is None:
        return f"{text} (client data)"
    mark = CONFIDENCE_MARK.get(cell.confidence, "")
    return f"{text} ({mark})" if mark else text


def build_document(run: BenchmarkRun, path: str | Path) -> str:
    docx = _require_docx()
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: PLC0415
    from docx.shared import Pt  # noqa: PLC0415

    doc = docx.Document()
    grid = build_cells(run)
    columns = [APEX_COLUMN] + list(run.results)

    doc.add_heading("Peer benchmarking — private credit comparables", level=0)
    sub = doc.add_paragraph()
    sub.add_run(
        f"Reporting quarter ended {run.anchor.isoformat()}. "
    ).bold = True
    sub.add_run(
        "Every figure is as reported by the filer for a period ending on or "
        "before that date, sourced directly from SEC EDGAR filings."
    )

    # ------------------------------------------------------------- the table
    doc.add_heading("Benchmark table", level=1)
    table = doc.add_table(rows=1, cols=len(columns) + 1)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    header[0].text = "Metric"
    for i, col in enumerate(columns, start=1):
        header[i].text = col
    for cells in table.rows[0].cells:
        for p in cells.paragraphs:
            for r in p.runs:
                r.bold = True

    for metric in ALL_METRICS:
        row = table.add_row().cells
        row[0].text = METRIC_LABELS.get(metric, metric)
        for i, col in enumerate(columns, start=1):
            row[i].text = _cell_text(grid[metric][col])

    note = doc.add_paragraph()
    note.add_run("Confidence: ").bold = True
    note.add_run(
        "High means several independent routes to the number agreed. Medium "
        "means corroborated but on weaker evidence. Low means a single "
        "uncorroborated source. A dash is not a missing number — it is a cell "
        "where the evidence did not support one, and the reason is stated."
    )

    # ------------------------------------------------------------- coverage
    rows = coverage_rows(run)
    filled = [r for r in rows if r.status == FILLED]
    doc.add_heading("Coverage", level=1)
    doc.add_paragraph(
        f"{len(filled)} of {len(rows)} competitor cells are populated. Each empty "
        "cell is attributed below so it can be read as a known limit rather than "
        "an unexplained gap."
    )
    buckets: dict[str, list] = {}
    for r in rows:
        if r.status != FILLED:
            buckets.setdefault(r.owner, []).append(r)
    for owner, label in (
        ("OURS", "Extraction we have not built"),
        ("CADENCE", "Exists, but outside the six-month staleness window"),
        ("CLIENT", "Withheld pending a client decision"),
        ("STRUCTURAL", "Not published by the filer; no work would fix these"),
    ):
        group = buckets.get(owner, [])
        if not group:
            continue
        doc.add_heading(f"{label} ({len(group)})", level=2)
        for r in sorted(group, key=lambda x: (x.fund, x.metric)):
            doc.add_paragraph(
                f"{r.fund} — {METRIC_LABELS.get(r.metric, r.metric)}: {r.detail}",
                style="List Bullet",
            )

    # ------------------------------------------------------------ conflicts
    conflicts = run.conflicts
    doc.add_heading("Source conflicts resolved", level=1)
    if not conflicts:
        doc.add_paragraph("No material source conflicts in this run.")
    else:
        doc.add_paragraph(
            "Where a filing supports more than one value for the same metric, "
            "both were extracted and the choice is recorded rather than made "
            "silently."
        )
        for ticker, rm in conflicts:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{ticker} — {METRIC_LABELS.get(rm.metric, rm.metric)}: ").bold = True
            p.add_run(
                f"candidates {', '.join(f'{v:g}' for v in rm.conflict.values)}; "
                f"resolved to {rm.conflict.resolution}. {rm.conflict.rationale}"
            )

    # ----------------------------------------------------------- comparison
    doc.add_heading("Apex Ridge versus the peer set", level=1)
    comps = compare(run)
    from ..config import APEX_BASIS_CONFIRMED  # noqa: PLC0415

    if not APEX_BASIS_CONFIRMED:
        warn = doc.add_paragraph()
        warn.add_run("Apex-versus-peer deltas are withheld. ").bold = True
        warn.add_run(
            "The share class and fee treatment behind Apex Ridge's own figures "
            "are not yet confirmed, and a delta between two numbers of unknown "
            "basis is precisely the confidently-wrong figure this process exists "
            "to prevent. Peer-to-peer statistics below do not depend on it."
        )
    ctable = doc.add_table(rows=1, cols=4)
    ctable.style = "Light Grid Accent 1"
    for i, h in enumerate(["Metric", "Peer median", "Peer range", "Apex Ridge"]):
        ctable.rows[0].cells[i].text = h
        for p in ctable.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for metric in ALL_METRICS:
        c = comps[metric]
        if not c.peers:
            continue
        unit = c_unit(metric)
        lo, hi = c.peer_range
        row = ctable.add_row().cells
        row[0].text = METRIC_LABELS.get(metric, metric)
        row[1].text = format_value(c.peer_median, unit)
        row[2].text = f"{format_value(lo, unit)} – {format_value(hi, unit)}"
        if c.apex is None:
            row[3].text = "—"
        elif not APEX_BASIS_CONFIRMED:
            row[3].text = f"{format_value(c.apex, unit)} (delta withheld)"
        else:
            bits = [format_value(c.apex, unit)]
            if c.apex_delta is not None:
                bits.append(f"{c.apex_delta:+.2f} vs median")
            if c.apex_rank:
                bits.append(f"rank {c.apex_rank[0]}/{c.apex_rank[1]}")
            row[3].text = ", ".join(bits)

    # ----------------------------------------------------------- provenance
    doc.add_heading("Provenance", level=1)
    doc.add_paragraph(
        "Every reported figure below cites the filing it came from. Accession "
        "numbers are permanent SEC identifiers and resolve years later, which is "
        "what a compliance review needs."
    )
    ptable = doc.add_table(rows=1, cols=4)
    ptable.style = "Light Grid Accent 1"
    for i, h in enumerate(["Fund", "Metric", "Source", "Location in document"]):
        ptable.rows[0].cells[i].text = h
        for p in ptable.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ticker, res in run.results.items():
        for metric in ALL_METRICS:
            rm = res.resolved.get(metric)
            if rm is None or rm.value is None or rm.chosen is None:
                continue
            prov = rm.chosen.provenance
            row = ptable.add_row().cells
            row[0].text = ticker
            row[1].text = METRIC_LABELS.get(metric, metric)
            row[2].text = (
                f"{prov.form_type}, period {prov.period_end}, acc. {prov.accession}"
            )
            row[3].text = prov.locator

    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.size is None:
                r.font.size = Pt(10)

    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)


def c_unit(metric: str) -> str:
    from ..config import METRIC_UNITS  # noqa: PLC0415

    return METRIC_UNITS.get(metric, "pct")
